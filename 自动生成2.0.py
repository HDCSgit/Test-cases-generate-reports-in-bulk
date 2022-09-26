# -*- coding:utf-8 -*-
# @Author : 郑凯繁
# @Time :2022/9/26 11:03

import os
import xlrd
from docx import Document


getPath = os.path.realpath(__file__)
rootPath = os.path.dirname(getPath)

#file_name = input('输入文件名称(必须xls格式且不包含后缀)')
file_name = 'SIT超网重建电子渠道项目案例v2.0'
data = xlrd.open_workbook(r'%s/%s.xls' % (rootPath, file_name))
#table = data.sheet_by_index(0)
sheet = data.sheet_names()

def creatDoc(nums, path):
    for i in range(len(nums)-1):
        document = Document()
        document.add_heading('%s' % case[i][0].strip('/'), 0)
        document.add_heading('一、用例编号-%s' % str(i+1).zfill(4))
        document.add_heading('二、测试目的：',level=1)
        document.add_paragraph('%s' % case[i][1])
        document.add_heading('三、前置条件：',level=1)
        document.add_paragraph('%s' % case[i][2])
        document.add_heading('四、实际操作步骤：',level=1)
        document.add_paragraph('%s' % case[i][3])
        document.add_heading('五、预期结果：',level=1)
        document.add_paragraph('%s' % case[i][4])
        document.add_heading('六、测试过程截图、登记簿查询、账户分录、关联系统检查、界面展示及回单凭证类检查等：\n',level=1)
        objpath_name = os.path.join(path,'%s-%s.docx' % (case[i][0].strip('/'),str(i+1).zfill(4)))
        document.save(objpath_name)
      
def newFile(rowNum,path):
    file_name = rowNum
    fpath =  path
    for name in file_name:
        os.mkdir(fpath+name)   
  
# for i in range(len(sheet)): # 设置遍历sheet范围 4的意思是遍历sheet1到sheet4,
# sheet名字可以自己随便取不受干扰
for i in range(4):
    #['所属模块', '用例标题', '前置条件', '步骤', '预期', '关键词', '优先级', '用例类型', '适用阶段'],
    table = data.sheet_by_name(sheet[i])
    
    rows_list = table.row_values(0)
    cols_list = table.col_values(0)
    
    # 新建文件夹，创建新文件路径，进入新路径
    os.mkdir(rootPath + './%s' % sheet[i])
    path = '%s/%s' % (rootPath, sheet[i])
    
    # 内容 case[行-1][列-1]
    case = []
    for i in range(1,table.nrows):
        case.append(table.row_values(i))
    
    doc = creatDoc(cols_list, path)
    #返回上一文件路径
    path = rootPath





    


